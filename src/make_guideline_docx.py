# -*- coding: utf-8 -*-
r"""계획지표 가이드라인 Word 생성 (config/02_계획지표_가이드라인.md 의 공유용 버전)."""
import os, sys, io, datetime
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PROJ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(PROJ, "outputs", datetime.date.today().strftime("%Y%m%d"),
                   "공원녹지계획_계획지표_가이드라인.docx")

def _kr(r): r.font.name="맑은 고딕"; r._element.rPr.rFonts.set(qn('w:eastAsia'),"맑은 고딕")
def setfont(d):
    st=d.styles["Normal"]; st.font.name="맑은 고딕"; st.font.size=Pt(10)
    st.element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'),"맑은 고딕")
def h(d,t,l=1):
    x=d.add_heading(t,level=l)
    for r in x.runs: _kr(r)
def para(d,t,size=10,bold=False,color=None):
    q=d.add_paragraph(); r=q.add_run(t); r.bold=bold; r.font.size=Pt(size); _kr(r)
    if color: r.font.color.rgb=color
def tbl(d,cols,rows,widths=None):
    t=d.add_table(rows=1,cols=len(cols)); t.style="Light Grid Accent 1"
    for j,c in enumerate(cols):
        cell=t.rows[0].cells[j]; cell.text=""; rr=cell.paragraphs[0].add_run(str(c))
        rr.bold=True; rr.font.size=Pt(8.5); _kr(rr); cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cs=t.add_row().cells
        for j,v in enumerate(row):
            cs[j].text=""; rr=cs[j].paragraphs[0].add_run(str(v)); rr.font.size=Pt(8.5); _kr(rr)
    if widths:
        for row in t.rows:
            for j,w in enumerate(widths): row.cells[j].width=Cm(w)

d=Document(); setfont(d)
s=d.sections[0]; s.top_margin=s.bottom_margin=Cm(2); s.left_margin=s.right_margin=Cm(2)
ti=d.add_heading("공원녹지계획 계획지표 적용 가이드라인\n— 녹지 접근성·연결성",level=0)
for r in ti.runs: _kr(r)
ti.alignment=WD_ALIGN_PARAGRAPH.CENTER
para(d,f"작성일 {datetime.date.today()} · 대상 대구광역시 · 근거 모듈4 분석(2019·2024)",9)

h(d,"0. 목적",1)
para(d,"모듈4 산출물인 '도보 도달시간'과 '전류밀도'를 공원녹지계획의 계획지표로 쓰기 위한 변환·기준·적용법을 제시한다. "
       "도달시간은 분 단위로 직관적이나, 전류밀도는 절대값이 무의미해 계획 언어(등급·용어)로 변환이 필요하다.")

h(d,"1. 두 계획지표 체계",1)
tbl(d,["구분","녹지 접근성","녹지 연결성"],
    [["의미","사람이 녹지를 얼마나 쉽게 누리나(공급·형평)","녹지가 얼마나 이어졌나(생태축·구조)"],
     ["원자료","도보 도달시간(분)","전류밀도(상대값)"],
     ["변환필요","낮음(직관적)","높음(등급·용어 변환)"],
     ["계획개념","생활권공원 공급, 녹지소외 해소","그린네트워크, 생태축 보전·복원"]])

h(d,"2. 지표1 — 접근성(도달시간 → 등급·기준)",1)
para(d,"2.1 도보시간 4등급 (국내외 기준 + 대구 실측 분포 반영)",10,bold=True)
tbl(d,["등급","도보시간","계획 용어","대구 인구비율(2024)"],
    [["Ⅰ 우수","≤5분(≈350m)","녹지 인접 생활권","49.5%"],
     ["Ⅱ 양호","5~10분(≈700m)","도보권 녹지 생활권","31.8%"],
     ["Ⅲ 보통","10~15분","준도보권","14.5%"],
     ["Ⅳ 취약","15분 초과","녹지소외지역","4.2%"]])
para(d,"2.2 기준 근거: 도보 10분 생활권(생활SOC), 근린공원 유치거리 500m, 3-30-300 원칙(300m), WHO 권고 "
       "→ '10분'을 목표선, '15분 초과'를 소외 기준선으로.",9)
para(d,"2.3 대표 KPI: 「도보 10분 내 녹지 도달 인구비율」 = 현황 81.3%. 보조: 녹지소외인구(15분 초과)=4.2%(약 10만명). "
       "목표 예시: 10분내 81.3%→90%(2030), 소외 4.2%→2%.",10)

h(d,"3. 지표2 — 연결성(전류밀도 → 계획 용어)",1)
para(d,"3.1 전류밀도는 절대값이 무의미(상대순위만 유효) → ①연결성지수(0~100, 백분위) 또는 ②등급·용어로 변환.",10)
para(d,"3.2 연결성 5등급 (대구 실측 백분위 기준) ★권장",10,bold=True)
tbl(d,["등급","백분위","대구 임계값","계획 용어","계획 취급"],
    [["1 핵심연결축","상위5%","≥0.62","광역 생태코리더/그린네트워크 골격","절대보전, 축폭 확보"],
     ["2 주요연결","5~20%","0.33~0.62","주요 녹지축","보전·완충"],
     ["3 일반연결","20~50%","0.15~0.33","보조 연결","유지·관리"],
     ["4 연결취약","하위30~50%","0.05~0.15","잠재 단절","징검돌녹지 보완"],
     ["5 연결단절","하위30%","<0.05","연결 단절지","복원·신규연결(그린인프라)"]])
para(d,"3.3 핀치포인트(병목): 고전류 + 좁은 목 = 끊기면 넓은 두 녹지 연결이 사라지는 지점. "
       "적은 투자로 큰 효과 → 최우선 보전·복원. 계획도서에 '우선관리 코리더'로 표기.",10,bold=True)

h(d,"4. 통합 — 우선순위 매트릭스",1)
tbl(d,["","연결성 낮음","연결성 높음"],
    [["접근성 낮음(소외)","★최우선: 신규공원 공급 + 연결복원","공급 확충(축은 양호)"],
     ["접근성 높음","연결 복원(공급은 충분)","유지·관리"]])
para(d,"접근성 등급도 × 연결성 등급도 → 4유형 처방도(주제도)로 지역별 처방 도출.",9)

h(d,"5. 계획 적용 5단계",1)
for t in ["① 현황진단 — 접근성·연결성 등급 주제도 + 지표표, 소외지역·단절지 목록화",
          "② 목표설정 — KPI 수치목표(10분 도달 90%, 소외 2%, 핀치포인트 N개소 보전)",
          "③ 전략수립 — 우선순위 매트릭스로 지역별 처방(공급/연결/유지)",
          "④ 실행 — 신규공원·소생활권녹지, 생태코리더 보전지구, 단절지 그린인프라",
          "⑤ 모니터링 — 본 분석 시스템 재실행으로 지표 변화 성과관리(반복형 설계)"]:
    para(d,t,10)
para(d,"계획도서 반영: 주제도 3종(접근성/연결성/통합처방), 지표표(전역·구군 현황·목표), 우선지역 목록.",9)

h(d,"6. 한계·유의",1)
para(d,"· 연결성 등급은 상대값 — 비교 시 동일 저항표·코어기준 유지, 값이 아닌 순위·등급으로 소통.",9)
para(d,"· 접근성은 도로망·인구 고정 → 변화는 녹지 공급 변화 중심 반영.",9)
para(d,"· 임계값은 대구 분포 기준 — 타 지역은 재설정. 전류밀도 상승≠개선(파편화 병목집중), R_eff와 함께 해석.",9)

d.save(OUT); print("저장:",OUT)
