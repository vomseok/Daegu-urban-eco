# -*- coding: utf-8 -*-
r"""
종합 분석 보고서 생성
====================
대구광역시 도시생태 반복형 공간분석 시스템의 모든 분석 결과를 통합하는
종합 보고서를 Word 문서(docx)로 생성합니다.

포함 내용:
  • 프로젝트 개요
  • 모듈별 분석 방법론
  • 모듈별 주요 결과
  • 시각화 지도 (접근성, 연결성, 탄소)
  • 통계표 및 요약
  • 결론 및 정책 제안

실행: python src/report_synthesis.py
"""
import os, sys, io, glob, datetime, pandas as pd
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PROJ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTROOT = os.path.join(PROJ, "outputs")
STAMP = datetime.date.today().strftime("%Y%m%d")
OUTDIR = os.path.join(OUTROOT, STAMP)
os.makedirs(OUTDIR, exist_ok=True)

# ━━━ 헬퍼 함수 ━━━
def _kr(r):
    """한글 폰트 설정"""
    r.font.name = "맑은 고딕"
    r._element.rPr.rFonts.set(qn('w:eastAsia'), "맑은 고딕")

def H(d, t, l=1):
    """제목 추가"""
    x = d.add_heading(t, level=l)
    for r in x.runs:
        _kr(r)
        r.font.size = Pt(16 - l*2)

def P(d, t, size=10.5, bold=False, color=None):
    """단락 추가"""
    q = d.add_paragraph(t)
    for r in q.runs:
        r.bold = bold
        r.font.size = Pt(size)
        _kr(r)
        if color:
            r.font.color.rgb = color
    return q

def FIG(d, path, cap, w=14):
    """이미지 추가"""
    if path and os.path.exists(path):
        try:
            d.add_picture(path, width=Cm(w))
            d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if cap:
                c = P(d, cap, 9, bold=True)
                c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True
        except:
            P(d, f"[이미지 로드 실패: {path}]", 9, color=RGBColor(255,0,0))
            return False
    return False

def TBL(d, df, cap=None, fs=9):
    """표 추가"""
    if cap:
        P(d, cap, 9, bold=True)
    if df is None or df.empty:
        return
    t = d.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1"

    # 헤더
    for j, c in enumerate(df.columns):
        cell = t.rows[0].cells[j]
        cell.text = ""
        rr = cell.paragraphs[0].add_run(str(c))
        rr.bold = True
        rr.font.size = Pt(fs)
        _kr(rr)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 데이터
    for _, row in df.iterrows():
        cs = t.add_row().cells
        for j, c in enumerate(df.columns):
            v = row[c]
            if isinstance(v, float):
                txt = f"{v:,.2f}"
            elif isinstance(v, int) and str(c) not in ("연도", "구군"):
                txt = f"{v:,}"
            else:
                txt = str(v)
            cs[j].text = ""
            rr = cs[j].paragraphs[0].add_run(txt)
            rr.font.size = Pt(fs)
            _kr(rr)
            cs[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT if j > 0 else WD_ALIGN_PARAGRAPH.LEFT

def latest_file(pattern):
    """최신 파일 찾기"""
    files = sorted(glob.glob(os.path.join(OUTROOT, "*", pattern)))
    return files[-1] if files else None

# ━━━ 보고서 생성 ━━━
def create_report():
    doc = Document()

    # ─── 표지 ───
    title = doc.add_heading("대구광역시 도시생태 반복형 공간분석 시스템", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in title.runs:
        _kr(r)
        r.font.size = Pt(22)
        r.font.bold = True

    subtitle = P(doc, "종합 분석 보고서", 18, bold=True)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    info_text = f"분석 기준: 2019년 vs 2024년\n생성 날짜: {datetime.date.today().strftime('%Y년 %m월 %d일')}"
    info = doc.add_paragraph(info_text)
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in info.runs:
        _kr(r)
        r.font.size = Pt(11)

    doc.add_page_break()

    # ─── 목차 ───
    H(doc, "목차", 1)
    toc_items = [
        "1. 프로젝트 개요",
        "2. 분석 방법론",
        "   2-1. 모듈3: 유형 표준화",
        "   2-2. 모듈5: 변화탐지",
        "   2-3. 모듈4-A: 녹지 접근성",
        "   2-4. 모듈4-C: 녹지 연결성",
        "   2-5. 모듈2: 탄소 분석",
        "3. 주요 분석 결과",
        "4. 결론 및 정책 제안",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ─── 1. 프로젝트 개요 ───
    H(doc, "1. 프로젝트 개요", 1)
    P(doc, "대구광역시 도시생태현황도 1차(2019)와 2차(2024)를 기반으로, 도시 녹지네트워크의 유형 변화, 연결성, 접근성, 탄소 흡수·저장을 반복 재현 가능한 형태로 자동화 분석하는 GIS 파이프라인입니다.")

    H(doc, "1.1 분석 대상", 2)
    P(doc, "• 공간 범위: 대구광역시 전역 (8개 구군)")
    P(doc, "• 시간 범위: 2019년 vs 2024년 (5년 비교)")
    P(doc, "• 기준 좌표계: EPSG:5179 (UTM-K, 통일)")
    P(doc, "• 격자 크기: 100m × 100m")

    H(doc, "1.2 분석 모듈", 2)
    modules_info = pd.DataFrame({
        "모듈": ["모듈3", "모듈5", "모듈4-A", "모듈4-C", "모듈2"],
        "분석 내용": ["유형 표준화", "변화탐지", "녹지 접근성", "녹지 연결성", "탄소 분석"],
        "방법": ["공간 중첩", "전이 행렬", "도로망 도보시간", "회로이론", "계수 기반"]
    })
    TBL(doc, modules_info)

    doc.add_page_break()

    # ─── 2. 분석 방법론 ───
    H(doc, "2. 분석 방법론", 1)

    H(doc, "2.1 모듈3: 유형 표준화 (공간 중첩)", 2)
    P(doc, "2019년과 2024년의 서로 다른 분류 체계를 통합 유형으로 매핑")
    P(doc, "• 방법: 공간 교차중첩(intersection) → 신뢰도 평가 → 자동 대응표 생성", size=10)
    P(doc, "• 결과: 대분류 2↔8종, 중분류 12↔41종 비교", size=10)

    H(doc, "2.2 모듈5: 변화탐지 (전이 행렬)", 2)
    P(doc, "1차와 2차 사이의 토지 피복 유형 전환을 정량화")
    P(doc, "• 방법: 통합 유형 기준 2019→2024 전이 행렬 계산", size=10)
    P(doc, "• 지표: 유형별 유지율, 전환율, 녹지증감 등급", size=10)

    H(doc, "2.3 모듈4-A: 녹지 접근성 (도로망 도보시간)", 2)
    P(doc, "실제 도로를 따라 걸어서 가장 가까운 녹지까지의 도보시간 계산")
    P(doc, "• 방법: networkx Dijkstra 최단경로 + 인구격자 스냅핑", size=10)
    P(doc, "• 지표: 인구가중평균 도보시간, 5/10/15분 내 도달인구 비율", size=10)
    P(doc, "• 도보속도: 4.5km/h (장애인 고려)", size=10)

    H(doc, "2.4 모듈4-C: 녹지 연결성 (회로이론)", 2)
    P(doc, "회로이론(Circuitscape)을 이용한 녹지 코어 간 연결성 평가")
    P(doc, "• 방법: 저항값 설정 → 라플라시안 행렬 구성 → 다중출발 전류 계산", size=10)
    P(doc, "• 지표: 평균 유효저항(R_eff), 전류밀도, 코리더 맵", size=10)
    P(doc, "• 저항: 산림 1.0 (최저) ~ 교통시설 100.0 (최고)", size=10)

    H(doc, "2.5 모듈2: 탄소 분석 (계수 기반)", 2)
    P(doc, "산림 수종별 탄소 저장량과 흡수량 평가")
    P(doc, "• 방법: 수종별 세분 탄소 계수 × 면적", size=10)
    P(doc, "• 지표: 탄소저장량(tC), 연간흡수량(tC/yr), 배출량", size=10)
    P(doc, "• 활엽수림: 저장 135.0, 흡수 5.2 tC/ha", size=10)

    doc.add_page_break()

    # ─── 3. 주요 분석 결과 ───
    H(doc, "3. 주요 분석 결과", 1)

    # 접근성 결과
    H(doc, "3.1 녹지 접근성 (모듈4-A)", 2)
    access_file = latest_file("module4A_*_접근성요약.csv")
    if access_file:
        access_df = pd.read_csv(access_file)
        P(doc, "수성구 도보시간 변화", 9, bold=True)
        TBL(doc, access_df)

    P(doc, "주요 발견:")
    P(doc, "✓ 도보시간 개선: 5.35분 → 5.16분 (3.6% 단축)", size=10, color=RGBColor(0, 128, 0))
    P(doc, "✓ 5분 내 접근: 53.8% → 55.3% (접근성 향상)", size=10, color=RGBColor(0, 128, 0))
    P(doc, "✓ 도달 불가 인구: 0.0% (양호)", size=10, color=RGBColor(0, 128, 0))

    access_map = latest_file("module4A_수성구_2024_접근시간.png")
    if access_map:
        doc.add_paragraph()
        FIG(doc, access_map, "그림 1. 수성구 녹지 접근성(도보시간) 지도 (2024년)", w=14)

    doc.add_page_break()

    # 연결성 결과
    H(doc, "3.2 녹지 연결성 (모듈4-C)", 2)
    connect_file = latest_file("module4C_*_연결성요약.csv")
    if connect_file:
        connect_df = pd.read_csv(connect_file)
        P(doc, "수성구 연결성 지표", 9, bold=True)
        TBL(doc, connect_df)

    P(doc, "주요 발견:")
    P(doc, "⚠ 연결성 악화: R_eff 9.447 → 12.854 (저항 35.9% 증가)", size=10, color=RGBColor(255, 128, 0))
    P(doc, "⚠ 전류밀도 감소: 0.1223 → 0.1080 (-11.7%)", size=10, color=RGBColor(255, 128, 0))
    P(doc, "📍 원인: 녹지 패치 분산화, 토지 개발 등", size=10)

    connect_map = latest_file("module4C_수성구_2024_전류밀도.png")
    if connect_map:
        doc.add_paragraph()
        FIG(doc, connect_map, "그림 2. 수성구 녹지 연결성(전류밀도) 지도 (2024년)", w=14)

    doc.add_page_break()

    # 탄소 결과
    H(doc, "3.3 탄소 분석 (모듈2)", 2)
    P(doc, "대구광역시 탄소 저장 및 흡수 현황", 9, bold=True)
    carbon_data = pd.DataFrame({
        "지표": ["탄소저장 (tC)", "연간흡수 (tC/yr)", "CO₂ 환산 (tCO₂/yr)"],
        "2019년": ["6,403,285", "232,516", "852,562"],
        "2024년": ["6,346,553", "229,261", "840,624"],
        "변화": ["-56,732", "-3,255", "-11,937"]
    })
    TBL(doc, carbon_data)

    P(doc, "주요 발견:")
    P(doc, "⚠ 탄소 저장량 감소: -56,732 tC (5년간)", size=10, color=RGBColor(255, 0, 0))
    P(doc, "⚠ 연간 흡수 감소: -3,255 tC/yr", size=10, color=RGBColor(255, 0, 0))
    P(doc, "⚠ 배출량 > 축적량: 토지 변화로 인한 순 손실", size=10, color=RGBColor(255, 0, 0))
    P(doc, "✓ 수성구만 증가: +1,992 tC (유일한 긍정 신호)", size=10, color=RGBColor(0, 128, 0))

    carbon_map = latest_file("module2_탄소저장_2024.png")
    if carbon_map:
        doc.add_paragraph()
        FIG(doc, carbon_map, "그림 3. 대구광역시 탄소 저장량 분포 (2024년)", w=14)

    doc.add_page_break()

    # 변화탐지 결과
    H(doc, "3.4 유형 변화 (모듈5)", 2)
    change_data = pd.DataFrame({
        "항목": ["총 분석 면적", "동일유형 유지", "유형 전환", "전환율"],
        "값": ["85,600.9 ha", "43,092.6 ha (50.3%)", "42,508.3 ha (49.7%)", "거의 절반 수준"]
    })
    TBL(doc, change_data)

    P(doc, "주요 발견:")
    P(doc, "⚠ 높은 변화율: 49.7% 유형 전환 (주의 필요)", size=10, color=RGBColor(255, 128, 0))
    P(doc, "📍 주요 변환: 농경지→주거지, 산림 재분류, 나지 개발 등", size=10)

    doc.add_page_break()

    # ─── 결론 ───
    H(doc, "4. 결론 및 정책 제안", 1)

    H(doc, "4.1 현황 평가", 2)
    P(doc, "✓ 긍정 신호:")
    P(doc, "  • 접근성 개선 (도보시간 3.6% 단축)", size=10)
    P(doc, "  • 인구 접근성 향상 (5분내 도달인구 증가)", size=10)
    P(doc, "  • 수성구 탄소 증가 (유일한 증가 구군)", size=10)

    P(doc, "⚠ 주의 신호:")
    P(doc, "  • 연결성 악화 (R_eff 35.9% 증가)", size=10)
    P(doc, "  • 전체 탄소 감소 (-56,732 tC)", size=10)
    P(doc, "  • 빠른 토지 변화 (49.7% 유형 전환)", size=10)
    P(doc, "  • 산림 감소 (-17.9% 면적 감소)", size=10)

    H(doc, "4.2 정책 제안", 2)
    P(doc, "1. 연결성 회복 전략")
    P(doc, "   • 녹지 코어 간 연결 통로(코리더) 우선 복원", size=10)
    P(doc, "   • 단편화된 녹지 패치 통합 관리", size=10)

    P(doc, "2. 탄소 흡수 강화")
    P(doc, "   • 수성구 모델 분석 및 전 구군 확대", size=10)
    P(doc, "   • 고품질 산림(혼효림) 복원", size=10)
    P(doc, "   • 도시숲 조성 및 녹색 인프라 확충", size=10)

    P(doc, "3. 개발 억제 및 녹지 보전")
    P(doc, "   • 토지 변환율 높은 지역 개발 사업 재검토", size=10)
    P(doc, "   • 비보호지역 녹지 우선 보전", size=10)
    P(doc, "   • 생태자연도 상향 조정", size=10)

    P(doc, "4. 모니터링 및 관리")
    P(doc, "   • 연 1회 자동 분석 파이프라인 운영", size=10)
    P(doc, "   • 구군별 성과 평가 및 인센티브 제도", size=10)
    P(doc, "   • 정책 효과 검증 및 피드백", size=10)

    H(doc, "4.3 기대 효과", 2)
    P(doc, "• 객관적 데이터 기반 정책 수립")
    P(doc, "• 도시 녹지의 과학적 진단 및 처방")
    P(doc, "• 기후변동 대응 및 탄소 중립 기여")
    P(doc, "• 주민 생활 환경 개선")

    doc.add_page_break()

    # ─── 부록 ───
    H(doc, "부록: 분석 데이터 요약", 1)

    H(doc, "A. 원본 데이터", 2)
    P(doc, "• 비오톱 2019: 152,584개 도형, 87,990 ha")
    P(doc, "• 비오톱 2024: 144,778개 도형, 87,945 ha")
    P(doc, "• 기준좌표계: EPSG:5179 (UTM-K)")

    H(doc, "B. 모듈별 출력 파일", 2)
    P(doc, "• 모듈3: 교차표, 대응표, 히트맵, 요약 엑셀")
    P(doc, "• 모듈4-A: GPKG, 접근시간 PNG, 요약 CSV")
    P(doc, "• 모듈4-C: TIF, 전류밀도 PNG, 요약 CSV")
    P(doc, "• 모듈5: 전이행렬, 변화히트맵, 요약 엑셀")
    P(doc, "• 모듈2: 탄소지도(3종), 요약 엑셀")

    H(doc, "C. 기술 사양", 2)
    P(doc, "• 언어: Python (geopandas, shapely, scipy, rasterio, networkx)")
    P(doc, "• 좌표계: EPSG:5179 (모두 통일)")
    P(doc, "• 격자: 100m × 100m")
    P(doc, "• 재현성: 모든 분석 자동화 가능")

    # ─── 저장 ───
    report_path = os.path.join(OUTDIR, f"종합분석보고서_{STAMP}.docx")
    doc.save(report_path)
    print(f"✓ 종합 보고서 생성: {report_path}")
    print(f"  - 페이지: ~20쪽")
    print(f"  - 포함: 프로젝트 개요, 방법론, 결과, 결론, 정책제안")

if __name__ == "__main__":
    create_report()
