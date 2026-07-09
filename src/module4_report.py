# -*- coding: utf-8 -*-
r"""모듈4 종합 보고서 생성 (python-docx) — 녹지 연결성·접근성, 2시점×3공간단위."""
import os, sys, io, datetime
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PROJ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
STAMP = datetime.date.today().strftime("%Y%m%d")
OUTDIR = os.path.join(PROJ, "outputs", STAMP)
XLSX = os.path.join(OUTDIR, "module4_3단위요약.xlsx")
def p(f): return os.path.join(OUTDIR, f)
def _kr(run): run.font.name="맑은 고딕"; run._element.rPr.rFonts.set(qn('w:eastAsia'),"맑은 고딕")

def setfont(doc):
    st=doc.styles["Normal"]; st.font.name="맑은 고딕"; st.font.size=Pt(10)
    st.element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'),"맑은 고딕")
def h(doc,t,l=1):
    x=doc.add_heading(t,level=l)
    for r in x.runs: _kr(r)
    return x
def para(doc,t,size=10,bold=False,color=None):
    q=doc.add_paragraph(); r=q.add_run(t); r.bold=bold; r.font.size=Pt(size); _kr(r)
    if color: r.font.color.rgb=color
    return q
def table(doc,df):
    t=doc.add_table(rows=1,cols=len(df.columns)); t.style="Light Grid Accent 1"
    for j,c in enumerate(df.columns):
        cell=t.rows[0].cells[j]; cell.text=""; rr=cell.paragraphs[0].add_run(str(c))
        rr.bold=True; rr.font.size=Pt(8.5); _kr(rr); cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    for _,row in df.iterrows():
        cells=t.add_row().cells
        for j,c in enumerate(df.columns):
            v=row[c]; txt=f"{v:,}" if isinstance(v,(int,float)) and c not in ("연도","구군") else str(v)
            cells[j].text=""; rr=cells[j].paragraphs[0].add_run(txt); rr.font.size=Pt(8.5); _kr(rr)
            cells[j].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT if j==0 else WD_ALIGN_PARAGRAPH.RIGHT
    return t
def img(doc,fname,width=15):
    if os.path.exists(p(fname)):
        doc.add_picture(p(fname),width=Cm(width))
        doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER

glob=pd.read_excel(XLSX,sheet_name="전역")
conn=pd.read_excel(XLSX,sheet_name="구군_연결성")
acc=pd.read_excel(XLSX,sheet_name="구군_접근성")

doc=Document(); setfont(doc)
s=doc.sections[0]; s.top_margin=s.bottom_margin=Cm(2); s.left_margin=s.right_margin=Cm(2)
ti=doc.add_heading("대구광역시 녹지 연결성·접근성 변화 분석\n(2019 → 2024, 전역·구군·100m격자)",level=0)
for r in ti.runs: _kr(r)
ti.alignment=WD_ALIGN_PARAGRAPH.CENTER
sub=para(doc,f"작성일 {datetime.date.today()} · 대상지 대구광역시 · 좌표계 EPSG:5179",9); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER

h(doc,"요약",1)
para(doc,"· 녹지(산림+조성녹지+수계·습지)의 연결성(회로이론)과 접근성(도로망 도보시간, 인구가중)을 "
        "2019·2024 두 시점에 대해 전역·구군·100m격자 3단위로 산출하였다.",10)
para(doc,"· 연결성은 5년간 악화되었다: 코어 간 유효저항 R_eff 18.56→22.34(+20%). "
        "녹지 단편화로 흐름이 좁은 핀치포인트에 집중되었다.",10)
para(doc,"· 접근성은 소폭 개선되었다: 인구가중 평균 도보 6.22→6.02분, 10분내 녹지 도달인구 80.1→81.3%. "
        "국지 녹지 면적 증가가 생활권 접근을 다소 높였다.",10)
para(doc,"· 즉 '가까운 녹지는 다소 늘었으나 큰 녹지축의 이어짐은 나빠진' 단편화가 핵심이다.",10)

h(doc,"1. 자료와 방법",1)
para(doc,"· 녹지: 통합비교유형의 산림·초지/조성녹지·수계/습지(연도별). 도로망: 도로명주소 도로구간(공식), 두 시점 공통.",10)
para(doc,"· 연결성(회로이론): 유형별 저항면을 100m 격자로 만들고, 주요 녹지코어(상위 10개, ≥50ha) 쌍마다 "
        "1A를 흘려 누적 전류밀도(코리더·핀치포인트)와 코어간 유효저항 R_eff을 계산.",10)
para(doc,"· 접근성: 도로구간을 교차점 노딩해 보행 그래프를 만들고(도보 4.5km/h), 녹지 진입노드 다중출발 "
        "최단경로로 각 100m 인구격자(국토통계, 인구가중)의 최근접 녹지 도보시간을 산출.",10)
para(doc,"· 3공간단위: 전역(대구 전체)·구군(8개)·100m격자(래스터/차이맵).",10)

h(doc,"2. 전역 결과",1)
para(doc,"표 2-1. 대구 전역 연결성·접근성 (2019 vs 2024)",9,bold=True)
table(doc,glob)
doc.add_paragraph()
para(doc,"그림 2-1. 연결성 변화 (2024-2019). 파랑=전류밀도 증가, 빨강=감소",9,bold=True)
img(doc,"module4_대구_연결성_차이.png",14)
para(doc,"그림 2-2. 접근성 변화 (2024-2019). 파랑=도보시간 단축(개선), 빨강=증가(악화)",9,bold=True)
img(doc,"module4_대구_접근성_차이.png",14)

h(doc,"3. 구·군 결과",1)
para(doc,"표 3-1. 구·군별 연결성(평균 전류밀도)",9,bold=True); table(doc,conn)
doc.add_paragraph()
para(doc,"표 3-2. 구·군별 접근성(인구가중 도보시간·10분내 도달%)",9,bold=True); table(doc,acc)
para(doc,"· 접근성 최악은 중구(11.4분)·서구(9.8분)로 도심 고밀지역이나, 서구는 5년간 -1.29분으로 개선폭이 가장 컸다.",10)
para(doc,"· 달서구·달성군은 접근성이 좋고(±) 추가 개선되었다. 북구는 접근성이 소폭 악화(+0.36분)되었다.",10)
para(doc,"· 연결성 전류밀도는 남부(달성군·달서구·남구)에서 증가, 북부 산지(동구·북구)에서 감소하는 공간 대비를 보였다.",10)
para(doc,"※ R_eff는 코어망 전체의 단일 스칼라라 구별로 분해되지 않으며, 구별 값은 '지역 녹지흐름이 그 구를 "
        "통과하는 강도(전류밀도)'로 해석한다. 전류밀도 증가가 곧 연결성 개선을 뜻하지는 않는다(핀치포인트 집중일 수 있음).",9)

h(doc,"4. 100m 격자 · 파일럿(수성구)",1)
para(doc,"· 100m 격자 산출물은 전역 래스터(tif)와 차이맵으로 제공된다(그림 2-1·2-2).",10)
para(doc,"그림 4-1. 수성구 확대 — 연결성(좌)·접근성(우) 코리더/도보시간",9,bold=True)
img(doc,"module4C_수성구_2024_전류밀도_경계.png",11)
img(doc,"module4A_수성구_2024_접근시간_경계.png",11)

h(doc,"5. 해석",1)
para(doc,"· 연결성 악화 + 접근성 국지개선의 병존은 '녹지 파편화'의 전형이다. 개별 소규모 녹지는 늘어 "
        "생활권 접근은 좋아졌으나, 대형 녹지코어 사이의 생태적 이어짐(회로 흐름)은 끊겨 R_eff가 상승했다.",10)
para(doc,"· 정책 함의: 접근성이 나쁜 도심(중구·서구·남구)의 소생활권 녹지 확충과, 연결성이 끊긴 북부 "
        "산지축의 코리더(핀치포인트) 보전·복원을 구분해 접근할 필요가 있다.",10)

h(doc,"6. 한계",1)
para(doc,"· 도로망·인구격자를 두 시점 공통(현재/2019)으로 고정 → 접근성 변화는 녹지 공급 변화만 반영.",9)
para(doc,"· 수계·습지를 녹지코어에 포함해 육상이동 저항은 근사이며, 저항값·코어기준은 config에서 조정 가능.",9)
para(doc,"· 연결성 저항면은 소분류 해상도이며 코어는 상위 10개 패치로 한정.",9)

out=p(f"녹지연결성_접근성_보고서_{STAMP}.docx"); doc.save(out)
print("저장:",out)
