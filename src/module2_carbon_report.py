# -*- coding: utf-8 -*-
r"""탄소 종합 보고서 생성 (python-docx). module2 산출물(엑셀·지도)을 최신 폴더에서 자동 탐색."""
import os, sys, io, glob, datetime
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PROJ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
def latest(pat):
    fs = sorted(glob.glob(os.path.join(PROJ, "outputs", "*", pat)))
    return fs[-1] if fs else None
XLSX   = latest("module2_탄소_3단위요약.xlsx")
M_SUM  = latest("module2_탄소계정_종합도.png")
M_STO  = latest("module2_탄소저장분포_2024.png")
M_BAL  = latest("module2_탄소수지_차이.png")
OUTDIR = os.path.dirname(M_SUM or XLSX)
STAMP  = datetime.date.today().strftime("%Y%m%d")

def _kr(r): r.font.name="맑은 고딕"; r._element.rPr.rFonts.set(qn('w:eastAsia'),"맑은 고딕")
def setfont(d):
    st=d.styles["Normal"]; st.font.name="맑은 고딕"; st.font.size=Pt(10)
    st.element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'),"맑은 고딕")
def h(d,t,l=1):
    x=d.add_heading(t,level=l)
    for r in x.runs: _kr(r)
def para(d,t,size=10,bold=False,color=None):
    q=d.add_paragraph(); r=q.add_run(t); r.bold=bold; r.font.size=Pt(size); _kr(r)
    if color: r.font.color.rgb=color
def table(d,df,fs=8.5):
    t=d.add_table(rows=1,cols=len(df.columns)); t.style="Light Grid Accent 1"
    for j,c in enumerate(df.columns):
        cell=t.rows[0].cells[j]; cell.text=""; rr=cell.paragraphs[0].add_run(str(c))
        rr.bold=True; rr.font.size=Pt(fs); _kr(rr); cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    for _,row in df.iterrows():
        cs=t.add_row().cells
        for j,c in enumerate(df.columns):
            v=row[c]; txt=f"{v:,.0f}" if isinstance(v,float) else (f"{v:,}" if isinstance(v,int) and str(c) not in("연도",) else str(v))
            cs[j].text=""; rr=cs[j].paragraphs[0].add_run(txt); rr.font.size=Pt(fs); _kr(rr)
            cs[j].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT if j==0 else WD_ALIGN_PARAGRAPH.RIGHT
def img(d,path,w=16):
    if path and os.path.exists(path):
        d.add_picture(path,width=Cm(w)); d.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER

glob_df=pd.read_excel(XLSX,sheet_name="전역")
gu_df=pd.read_excel(XLSX,sheet_name="구군")
sp=pd.read_excel(XLSX,sheet_name="유형별상세")
cf=pd.read_excel(XLSX,sheet_name="계수표")
forest=sp[sp["탄소유형"].isin(["침엽수림","활엽수림","혼효림","인공림","산림기타"])].copy()

d=Document(); setfont(d)
s=d.sections[0]; s.top_margin=s.bottom_margin=Cm(2); s.left_margin=s.right_margin=Cm(2)
ti=d.add_heading("대구광역시 탄소계정 분석 보고서\n(도시생태현황도 기반, 2019·2024)",level=0)
for r in ti.runs: _kr(r)
ti.alignment=WD_ALIGN_PARAGRAPH.CENTER
para(d,f"작성일 {datetime.date.today()} · 대상 대구광역시 · 좌표계 EPSG:5179 · 격자 100m",9)

h(d,"요약",1)
para(d,"· 도시생태현황도 유형별 탄소저장(tC/ha)·연간흡수(tC/ha/yr) 계수를 결합해 대구 전역·구군·100m격자로 "
       "탄소계정을 산정하고, 2019→2024 저장변화로 토지변화 배출·축적을 계산하였다.")
para(d,"· 산림 탄소는 수종(침엽·활엽·혼효)으로 세분해 그 해의 탄소 분포를 정밀 표현하였다.")
para(d,f"· 대구 전역 탄소저장은 약 {glob_df.iloc[0]['2024']:,.0f} tC이며, 연간 흡수는 약 "
       f"{glob_df.iloc[1]['2024']:,.0f} tC/yr(≈{glob_df.iloc[2]['2024']:,.0f} tCO₂/yr)이다.")
para(d,"· 5년간 토지변화로 순 탄소저장이 감소하였다(녹지→시가화 전용에 따른 배출이 축적을 상회).",
     color=RGBColor(0xB0,0x30,0x00))

h(d,"1. 자료와 방법",1)
para(d,"1.1 계수 출처",10,bold=True)
para(d,"· 산림·습지 저장/흡수: 국립산림과학원(2022)[원고 표4-13] — 참나무류림 135 tC/ha·5.2, 소나무림 98·4.1, 혼효림 118·4.8.")
para(d,"· 도시림 흡수 2.4 tC/ha/yr: Nowak et al.(2013). 토양탄소(초지·농경지·시가화·나지): IPCC 2006 Tier1 + 이경학 외(2006) + 박은진·강규이(2009).")
para(d,"1.2 유형 결합·수종 세분",10,bold=True)
para(d,"· 2019는 U소분류(산림은 U세분류로 수종 판별), 2024는 유형중분류로 통합비교유형에 결합. "
       "산림은 침엽·활엽·혼효·인공·기타로 세분해 계수를 차등 적용.")
para(d,"1.3 변화·배출 산정(중요)",10,bold=True)
para(d,"· 2019 U세분류와 2024 유형중분류의 산림 분류체계가 달라(2019 혼효림 다수 ↔ 2024 침엽·인공 위주) "
       "수종 세분값을 그대로 연도 비교하면 재분류 아티팩트가 발생한다. 따라서 연도 비교(배출·축적)는 "
       "산림을 단일계수(저장 115 tC/ha·흡수 4.6)로 평탄화해 비교 가능하게 산정하였다. 수종 세분값은 각 연도 현황(스냅샷)에만 사용한다.")

h(d,"2. 전역 탄소계정",1)
para(d,"표 2-1. 대구 전역 탄소저장·흡수 (2019 vs 2024)",9,bold=True)
table(d,glob_df)
emit=int(gu_df["배출_tC"].sum()); gain=int(gu_df["축적_tC"].sum())
para(d,f"· 토지변화 배출 {emit:,} tC(≈{emit*44/12:,.0f} tCO₂), 축적 {gain:,} tC → 순 {gain-emit:,} tC(손실).",10,bold=True)
para(d,"그림 2-1. 대구 탄소계정 공간표현 (A 저장 · B 흡수 · C 토지변화 수지 · D 구별 순배출)",9,bold=True)
img(d,M_SUM,16.5)

h(d,"3. 수종별 탄소 (현황 정밀)",1)
para(d,"표 3-1. 산림 수종별 면적·탄소저장 (스냅샷; 2019·2024 산림분류체계 상이로 직접비교 주의)",9,bold=True)
table(d,forest)
para(d,"※ 2019 '혼효림'의 상당수가 2024 유형중분류에서 침엽·활엽·인공림으로 재배정된다. 이는 실제 산림변화가 "
       "아니라 분류체계 차이이므로, 위 표는 각 연도 내 구성 파악용이며 연도 간 수종 증감으로 해석하지 않는다.",9)
para(d,"그림 3-1. 탄소저장 밀도 분포 2024 (수종 세분 반영, tC/ha)",9,bold=True)
img(d,M_STO,13)

h(d,"4. 구·군 탄소계정",1)
para(d,"표 4-1. 구·군별 탄소저장·흡수·토지변화 수지",9,bold=True)
table(d,gu_df)
top=gu_df.sort_values("배출_tC",ascending=False).iloc[0]
para(d,f"· 배출 최대는 {top['구군']}({int(top['배출_tC']):,} tC)로 산지 전용에 기인한다. 달성군·북구·동구가 "
       "주요 배출원, 수성구·달서구 등은 순 축적(음의 배출)을 보인다.",10)

h(d,"5. 토지변화 탄소수지",1)
para(d,"그림 5-1. 토지변화 탄소수지 (2019→2024). 빨강=배출(녹지→시가화 저장손실), 파랑=축적",9,bold=True)
img(d,M_BAL,13)
para(d,"· 배출 집중지는 북·동구 산지 전용 지점이며, 하천·복원지에서 소규모 축적이 나타난다. 앞선 연결성 악화·"
       "개발 사이클과 일관된 결과이다.",10)

h(d,"6. 한계",1)
para(d,"· 비산림(초지·농경·시가화·나지) 계수는 문헌 근사치로, 절대량보다 상대·변화 방향의 신뢰도가 높다. "
       "config/carbon_coefficients.csv 에서 조정 후 재실행 가능.",9)
para(d,"· 배출은 '토지변화에 따른 저장탄소 손실'로 산정하며, 화석연료·에너지 배출은 포함하지 않는다.",9)
para(d,"· 산림 수종 세분은 두 조사의 분류체계 차이로 연도 비교엔 부적합(평탄화로 보정).",9)

h(d,"부록. 탄소계수표",1)
table(d,cf[["탄소유형","탄소저장_tC_ha","연간흡수_tC_ha_yr","근거_비고"]],fs=8)

out=os.path.join(OUTDIR,f"대구_탄소계정_보고서_{STAMP}.docx"); d.save(out)
print("저장:",out)
