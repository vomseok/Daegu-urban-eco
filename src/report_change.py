# -*- coding: utf-8 -*-
r"""
변화탐지 종합 보고서 생성기 (python-docx)
==========================================
모듈5-R(유형변화)·모듈5-G(등급변화) 산출물을 읽어 Word 보고서를 만든다.
지도(히트맵)·표·자동서술·방법론·한계 포함.

입력:  outputs\<실행일자>\module5R_변화요약.csv, module5R_전이히트맵.png
       outputs\<실행일자>\module5G_등급요약.csv, module5G_등급히트맵.png
출력:  outputs\<실행일자>\변화탐지_종합보고서_YYYYMMDD.docx
실행:  .venv\Scripts\python.exe src\report_change.py
"""
import os, sys, io, datetime
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

PROJ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
STAMP = datetime.date.today().strftime("%Y%m%d")
OUTDIR = os.path.join(PROJ, "outputs", STAMP)

def p(path): return os.path.join(OUTDIR, path)

# ── 스타일 헬퍼 ──────────────────────────────────────────────────────
def set_font(doc, name="맑은 고딕", size=10):
    st = doc.styles["Normal"]
    st.font.name = name; st.font.size = Pt(size)
    from docx.oxml.ns import qn
    st.element.rcPar = None
    rpr = st.element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    rf.set(qn('w:eastAsia'), name)

def h(doc, text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs:
        r.font.name = "맑은 고딕"
        from docx.oxml.ns import qn
        r._element.rPr.rFonts.set(qn('w:eastAsia'), "맑은 고딕")
    return hd

def para(doc, text, size=10, bold=False, italic=False, color=None):
    q = doc.add_paragraph()
    r = q.add_run(text); r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.name = "맑은 고딕"
    from docx.oxml.ns import qn
    r._element.rPr.rFonts.set(qn('w:eastAsia'), "맑은 고딕")
    if color: r.font.color.rgb = color
    return q

def table_from_df(doc, df, widths=None, headshade="D9E1F2"):
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    hdr = t.rows[0].cells
    for j, c in enumerate(df.columns):
        hdr[j].text = ""
        rr = hdr[j].paragraphs[0].add_run(str(c)); rr.bold = True; rr.font.size = Pt(9)
        rr.font.name = "맑은 고딕"; rr._element.rPr.rFonts.set(qn('w:eastAsia'), "맑은 고딕")
        hdr[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(df.columns):
            v = row[c]
            txt = f"{v:,}" if isinstance(v, (int, float)) and c != "등급" else str(v)
            cells[j].text = ""
            rr = cells[j].paragraphs[0].add_run(txt); rr.font.size = Pt(9)
            rr.font.name = "맑은 고딕"; rr._element.rPr.rFonts.set(qn('w:eastAsia'), "맑은 고딕")
            cells[j].paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.RIGHT)
    return t

# ── 문서 작성 ────────────────────────────────────────────────────────
doc = Document()
set_font(doc)
sec = doc.sections[0]
sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)

# 표지
title = doc.add_heading("대구광역시 도시생태현황도\n1차(2019) ↔ 2차(2024) 변화탐지 보고서", level=0)
for r in title.runs:
    from docx.oxml.ns import qn
    r.font.name = "맑은 고딕"; r._element.rPr.rFonts.set(qn('w:eastAsia'), "맑은 고딕")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = para(doc, f"작성일: {datetime.date.today()}  ·  대상지: 대구광역시  ·  기준좌표계: EPSG:5179",
           size=10, italic=True)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# 요약
h(doc, "요약 (Executive Summary)", 1)
para(doc, "· 2019·2024 두 시점의 비오톱 도형(15.3만/14.5만)을 EPSG:5179로 통일하고, 분류체계 차이를 "
          "공간중첩 기반 대응표로 해소한 뒤 유형·등급 변화를 산출하였다.", 10)
para(doc, "· 유형 변화(통합 14유형 기준): 전체의 93.7%가 유형을 유지, 6.3%(약 5,552 ha)가 전환. "
          "변화의 핵심은 개발대기지(나지·유휴지)의 3배 증가로, 초지·조성녹지·공업지·농경지에서 유입되어 "
          "일부가 공동주택지로 전환되는 도시개발 사이클이 확인된다.", 10)
para(doc, "· 산림(99.0% 유지)·수계(96.7%) 등 자연기반은 대체로 안정적이다.", 10)
para(doc, "· 등급 변화는 두 조사의 채점 체계가 달라(2019 2등급 ≈ 2024 1등급) 직접 비교가 어렵다. 참고자료로만 제시한다.", 10)

# 1. 방법
h(doc, "1. 자료와 방법", 1)
para(doc, "1.1 자료", 10, bold=True)
para(doc, "· 1차(2019): 대구시 비오톱 등급평가 결과 (152,584 도형, 원본 EPSG:5186)\n"
          "· 2차(2024~25): 제2차 도시생태현황지도 보완제출용 (144,778 도형, 원본 EPSG:5187)", 10)
para(doc, "1.2 전처리", 10, bold=True)
para(doc, "· 좌표계 EPSG:5179로 통일, 깨진 도형 복구(1차 237·2차 767건), 속성 라벨 정제.", 10)
para(doc, "1.3 분류체계 표준화(핵심)", 10, bold=True)
para(doc, "· 1차(U대/중/소/세분류)와 2차(유형/피복/이용) 체계가 달라 직접 비교가 불가하다. "
          "공간중첩(intersection) 면적을 기준으로 대응관계를 도출하고, 자연계열의 왜곡을 막기 위해 "
          "2019는 U소분류, 2024는 유형중분류를 각각 14개 '통합비교유형'으로 재라벨하여 동일 해상도로 비교하였다.", 10)
para(doc, "· 유형별로 융합(dissolve) 후 두 시점을 교차중첩하여 전이면적(ha) 행렬을 산출하였다.", 10)

# 2. 유형 변화
h(doc, "2. 토지유형 변화 (2019 → 2024)", 1)
r = pd.read_csv(p("module5R_변화요약.csv"))
para(doc, "표 2-1. 통합비교유형별 면적·유지·순변화 (단위: ha)", 9, bold=True)
table_from_df(doc, r.rename(columns={"통합비교유형":"유형"}))
doc.add_paragraph()
if os.path.exists(p("module5R_전이히트맵.png")):
    para(doc, "그림 2-1. 유형 전이 히트맵 (행=2019, 열=2024; 대각선=유지)", 9, bold=True)
    doc.add_picture(p("module5R_전이히트맵.png"), width=Cm(15))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

para(doc, "2.1 주요 해석", 10, bold=True)
kept_r = 93.7
para(doc, f"· 전체 유지율 {kept_r}% — 5년 간 대부분의 토지유형은 그대로 유지되었다.", 10)
para(doc, "· 개발 사이클: 나지·유휴지가 863→2,586 ha로 3배 증가(+1,723). 유입원은 초지·조성녹지(-842), "
          "공업지(-617), 농경지(-442)이며, 나지·유휴지의 일부(133 ha)는 공동주택지로 전환되었다. "
          "공동주택지는 +227 ha 증가, 단독주택지는 소폭 감소하여 아파트화 경향을 보인다.", 10)
para(doc, "· 자연기반 안정: 산림 99.0%, 수계·습지 96.7%, 농경지 93.9% 유지. 산림 순변화는 -346 ha로 미미하다.", 10)
para(doc, "· 초지·조성녹지는 유지율 47.4%로 가장 유동적이며 대부분 개발대기지로 전환되었다 — 도시녹지 관리의 주목 지점.", 10)

# 3. 등급 변화(참고)
h(doc, "3. 비오톱 등급 변화 (참고)", 1)
g = pd.read_csv(p("module5G_등급요약.csv"))
g["등급"] = g["등급"].astype(str)
para(doc, "표 3-1. 등급별 면적·유지 (2024 최종평가등은 앞자리 주등급으로 정규화; 단위 ha)", 9, bold=True)
table_from_df(doc, g)
doc.add_paragraph()
if os.path.exists(p("module5G_등급히트맵.png")):
    para(doc, "그림 3-1. 등급 전이 히트맵", 9, bold=True)
    doc.add_picture(p("module5G_등급히트맵.png"), width=Cm(12))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
para(doc, "⚠ 해석 주의: 1등급이 19,644→54,167 ha로 급증하고 2019 2등급의 대부분(20,309 ha)이 2024 1등급으로 "
          "이동한 것은, 생태 개선이라기보다 두 조사의 등급 채점 기준이 다름을 시사한다(2019 2등급 ≈ 2024 1등급). "
          "따라서 등급 전이는 '체계 차이'를 걷어낸 뒤에만 변화로 해석해야 하며, 본 절은 참고자료이다.",
     10, color=RGBColor(0xB0, 0x30, 0x00))

# 4. 한계·다음단계
h(doc, "4. 한계와 다음 단계", 1)
para(doc, "· 자연계열은 소분류 해상도로 통합했으므로(산림 1종 등) 수종·천이 수준의 변화는 포착하지 않는다. "
          "필요 시 U세분류 기반 세분화가 가능하다.", 10)
para(doc, "· 도로녹지↔교통, 물류시설지↔공공용도 등 일부 경계 유형은 매핑 판단에 따라 결과가 달라질 수 있다"
          "(config/crosswalk/통합비교유형_*.csv 에서 조정 가능).", 10)
para(doc, "· <1% 규모의 산림↔수계 등 미세 교차는 두 조사의 경계·판독 차이(노이즈)를 포함한다.", 10)
para(doc, "· 다음 단계: 등급체계 정합(대응표) 확정 → 모듈4(녹지·생태축 연결성), 모듈6(보전·개발 우선지역)으로 확장.", 10)

doc.add_paragraph()
foot = para(doc, "본 보고서는 자동 생성되었습니다 (src/report_change.py). 원자료: outputs/%s/" % STAMP, 8, italic=True)
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

outp = p(f"변화탐지_종합보고서_{STAMP}.docx")
doc.save(outp)
print("저장:", outp)
